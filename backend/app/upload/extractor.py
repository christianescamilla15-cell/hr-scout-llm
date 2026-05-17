"""CV file extraction — PDF + DOCX → plain text + heuristic PII.

Per spec §4 + §8:
- We never persist the original file. We extract text, optionally
  pull a heuristic `full_name` + `email`, and the orig bytes are
  garbage-collected as soon as this function returns.
- Max 10 MB per file (enforced by the router, not here).
- PDFs may have columns/tables; pdfplumber handles both better than
  pypdf (slower, but accuracy matters when scoring on the text).
"""

import io
import logging
import re
from dataclasses import dataclass
from typing import Literal

import pdfplumber
from docx import Document

log = logging.getLogger(__name__)


class ExtractionError(Exception):
    """Raised when the file is unreadable / corrupt / unsupported."""


ALLOWED_MIME_TYPES: dict[str, Literal["pdf", "docx"]] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}

ALLOWED_EXTENSIONS: dict[str, Literal["pdf", "docx"]] = {
    ".pdf": "pdf",
    ".docx": "docx",
}

# RFC 5322 simplified — good enough for CVs in practice. Avoids the full
# 200-line monster and false-rejects exotic-but-valid edge cases like
# `+tag@`.
_EMAIL_RE = re.compile(
    r"(?:[A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,})", re.IGNORECASE
)

# Heuristic name lines: first 5 non-blank lines, scored on:
#   - 2-4 capitalized words (most CVs put the name at the very top)
#   - no digits, no email-shape, no URL-shape
_NAME_LINE_RE = re.compile(r"^[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+){1,3}\s*$")


@dataclass
class ExtractedCV:
    text: str
    full_name: str | None
    email: str | None
    pages: int  # 1 for DOCX (no pagination concept here)
    source: Literal["pdf", "docx"]


def detect_kind(filename: str, content_type: str | None) -> Literal["pdf", "docx"]:
    """Trust the filename extension first (more reliable than browser-set
    content-type), fall back to MIME. Raise if neither matches."""
    if filename:
        ext = filename.lower().rsplit(".", 1)
        if len(ext) == 2 and "." + ext[1] in ALLOWED_EXTENSIONS:
            return ALLOWED_EXTENSIONS["." + ext[1]]
    if content_type and content_type in ALLOWED_MIME_TYPES:
        return ALLOWED_MIME_TYPES[content_type]
    raise ExtractionError(
        f"Unsupported file type. Filename: {filename!r}, mime: {content_type!r}. "
        "Allowed: .pdf, .docx"
    )


def _try_extract_name(text: str) -> str | None:
    """Pick the first line in the top 5 non-blank lines that looks like a name."""
    seen_lines = 0
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        seen_lines += 1
        if seen_lines > 5:
            break
        if _EMAIL_RE.search(stripped):
            continue
        if any(c.isdigit() for c in stripped):
            continue
        if stripped.startswith("http") or "/" in stripped:
            continue
        if _NAME_LINE_RE.match(stripped) and 2 <= len(stripped.split()) <= 4:
            return stripped[:255]
    return None


def _try_extract_email(text: str) -> str | None:
    m = _EMAIL_RE.search(text)
    return m.group(0).lower()[:320] if m else None


def extract_pdf(data: bytes) -> ExtractedCV:
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = pdf.pages
            chunks: list[str] = []
            for page in pages:
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append(text)
            combined = "\n\n".join(chunks).strip()
            page_count = len(pages)
    except Exception as exc:
        raise ExtractionError(f"PDF parse failed: {exc}") from exc

    if not combined:
        raise ExtractionError(
            "PDF contains no extractable text (likely a scanned image — try DOCX or pasted text)"
        )

    return ExtractedCV(
        text=combined,
        full_name=_try_extract_name(combined),
        email=_try_extract_email(combined),
        pages=page_count,
        source="pdf",
    )


def extract_docx(data: bytes) -> ExtractedCV:
    try:
        doc = Document(io.BytesIO(data))
        lines: list[str] = []
        for paragraph in doc.paragraphs:
            t = paragraph.text.strip()
            if t:
                lines.append(t)
        # Tables (some CVs put contact info there)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in lines:
                        lines.append(t)
    except Exception as exc:
        raise ExtractionError(f"DOCX parse failed: {exc}") from exc

    combined = "\n".join(lines).strip()
    if not combined:
        raise ExtractionError("DOCX contains no extractable text")

    return ExtractedCV(
        text=combined,
        full_name=_try_extract_name(combined),
        email=_try_extract_email(combined),
        pages=1,
        source="docx",
    )


def extract(data: bytes, filename: str, content_type: str | None) -> ExtractedCV:
    kind = detect_kind(filename, content_type)
    if kind == "pdf":
        return extract_pdf(data)
    return extract_docx(data)
